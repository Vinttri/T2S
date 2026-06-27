"""Arbitrary-document → plain-text extraction for the agent schema loader.

The agent loader does NOT assume YAML structure. Enrichment documents can be any
business artifact: a data dictionary in Markdown, a CSV of column definitions, a
PDF spec, a Word glossary, an Excel sheet of code mappings, a JSON export, etc.
``extract_text(filename, data)`` dispatches purely on the file extension and
returns plain UTF-8 text that downstream LLM enrichment reads.

Design rules:
  * Text-ish formats (.md/.txt/.csv/.tsv/.json/.yml/.yaml/.sql/.log/.rst/...)
    decode directly — no third-party dependency.
  * Binary formats degrade gracefully: ``.pdf`` via ``pypdf``, ``.docx`` via
    ``python-docx``, ``.xlsx``/``.xlsm`` via ``openpyxl``. Each importer sits
    behind its own try/except so a MISSING optional library produces a short
    "unsupported, skipped" marker instead of crashing the whole enrichment run.
  * Nothing here ever raises to the caller: a per-file failure returns a marker
    string so one bad upload cannot abort the others.
"""

import csv
import io
import logging
from typing import Callable, Dict

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# Extensions decoded directly as UTF-8 text (no external library needed).
_TEXT_EXTENSIONS = {
    "md", "markdown", "txt", "text", "csv", "tsv", "json", "yml", "yaml",
    "sql", "log", "rst", "ini", "cfg", "conf", "toml", "tab", "ndjson",
    "html", "htm", "xml",
}

_MAX_BYTES_DEFAULT = 5_000_000  # guard a single uploaded file from blowing memory


def _decode_bytes(data: bytes) -> str:
    """Best-effort UTF-8 decode with a latin-1 fallback (never raises)."""
    if isinstance(data, str):
        return data
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, AttributeError):
            continue
    return data.decode("utf-8", errors="replace")


def _extract_text_plain(filename: str, data: bytes) -> str:
    """Decode plain-text / structured-text uploads verbatim."""
    return _decode_bytes(data)


def _extract_text_csv(filename: str, data: bytes) -> str:
    """Render CSV/TSV as compact ``col | col | col`` lines for the LLM.

    Falls back to a raw decode if the dialect cannot be sniffed.
    """
    text = _decode_bytes(data)
    delimiter = "\t" if filename.lower().endswith((".tsv", ".tab")) else ","
    try:
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        lines = [" | ".join(str(cell).strip() for cell in row) for row in reader]
        rendered = "\n".join(line for line in lines if line.strip())
        return rendered or text
    except Exception as exc:  # pylint: disable=broad-exception-caught
        logging.info("doc_text: CSV parse fell back to raw decode for %s: %s", filename, exc)
        return text


def _extract_text_pdf(filename: str, data: bytes) -> str:
    """Extract text from a PDF via ``pypdf`` (optional dependency)."""
    try:
        from pypdf import PdfReader  # pylint: disable=import-outside-toplevel
    except Exception as exc:  # pylint: disable=broad-exception-caught
        logging.info("doc_text: pypdf unavailable, skipping %s: %s", filename, exc)
        return f"[unsupported file (pypdf not installed), skipped: {filename}]"

    try:
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:  # pylint: disable=broad-exception-caught
                continue
        return "\n".join(part for part in parts if part.strip())
    except Exception as exc:  # pylint: disable=broad-exception-caught
        logging.info("doc_text: failed to read PDF %s: %s", filename, exc)
        return f"[failed to read PDF, skipped: {filename}]"


def _extract_text_docx(filename: str, data: bytes) -> str:
    """Extract paragraph + table text from a .docx via ``python-docx``."""
    try:
        import docx  # pylint: disable=import-outside-toplevel
    except Exception as exc:  # pylint: disable=broad-exception-caught
        logging.info("doc_text: python-docx unavailable, skipping %s: %s", filename, exc)
        return f"[unsupported file (python-docx not installed), skipped: {filename}]"

    try:
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join(cell for cell in cells if cell)
                if line:
                    parts.append(line)
        return "\n".join(parts)
    except Exception as exc:  # pylint: disable=broad-exception-caught
        logging.info("doc_text: failed to read DOCX %s: %s", filename, exc)
        return f"[failed to read DOCX, skipped: {filename}]"


def _extract_text_xlsx(filename: str, data: bytes) -> str:
    """Render each .xlsx sheet as ``Sheet: name`` + pipe-joined rows via openpyxl."""
    try:
        from openpyxl import load_workbook  # pylint: disable=import-outside-toplevel
    except Exception as exc:  # pylint: disable=broad-exception-caught
        logging.info("doc_text: openpyxl unavailable, skipping %s: %s", filename, exc)
        return f"[unsupported file (openpyxl not installed), skipped: {filename}]"

    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in workbook.worksheets:
            parts.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        try:
            workbook.close()
        except Exception:  # pylint: disable=broad-exception-caught
            pass
        return "\n".join(parts)
    except Exception as exc:  # pylint: disable=broad-exception-caught
        logging.info("doc_text: failed to read XLSX %s: %s", filename, exc)
        return f"[failed to read XLSX, skipped: {filename}]"


# Extension → extractor dispatch table.
_DISPATCH: Dict[str, Callable[[str, bytes], str]] = {
    "csv": _extract_text_csv,
    "tsv": _extract_text_csv,
    "tab": _extract_text_csv,
    "pdf": _extract_text_pdf,
    "docx": _extract_text_docx,
    "xlsx": _extract_text_xlsx,
    "xlsm": _extract_text_xlsx,
}


def _extension(filename: str) -> str:
    name = str(filename or "").strip().lower()
    if "." not in name:
        return ""
    return name.rsplit(".", 1)[-1]


def extract_text(filename: str, data: bytes, max_bytes: int = _MAX_BYTES_DEFAULT) -> str:
    """Return plain text for an uploaded document, dispatched by extension.

    Never raises. Unknown or unreadable formats yield a short marker string so
    the caller can log and skip them while continuing with the rest.

    Args:
        filename: Original upload name (used only for the extension).
        data: Raw file bytes (``str`` is also accepted for already-decoded text).
        max_bytes: Hard cap on bytes read from a single file.
    """
    name = str(filename or "metadata")
    raw = data if isinstance(data, (bytes, bytearray)) else str(data).encode("utf-8")
    if isinstance(raw, (bytes, bytearray)) and max_bytes > 0 and len(raw) > max_bytes:
        logging.info(
            "doc_text: truncating %s from %d to %d bytes",
            name, len(raw), max_bytes,
        )
        raw = bytes(raw[:max_bytes])

    ext = _extension(name)
    try:
        if ext in _DISPATCH:
            return _DISPATCH[ext](name, bytes(raw))
        if ext in _TEXT_EXTENSIONS or ext == "":
            return _extract_text_plain(name, bytes(raw))
        # Unknown binary: try a tolerant decode; if it looks like noise, skip.
        decoded = _decode_bytes(bytes(raw))
        printable = sum(1 for ch in decoded[:2000] if ch.isprintable() or ch in "\r\n\t")
        if decoded and printable >= 0.8 * min(len(decoded), 2000):
            return decoded
        return f"[unsupported file type '.{ext}', skipped: {name}]"
    except Exception as exc:  # pylint: disable=broad-exception-caught
        logging.info("doc_text: extraction failed for %s: %s", name, exc)
        return f"[failed to extract text, skipped: {name}]"


def extract_documents(
    documents,
    max_total_chars: int = 400_000,
    per_file_header: bool = True,
) -> str:
    """Concatenate extracted text from many ``(filename, bytes)`` documents.

    Each file is prefixed with a ``===== FILE: name =====`` header so the LLM can
    attribute facts to a source. The concatenation is bounded by
    ``max_total_chars`` to keep the prompt within limits; truncation is marked.
    """
    parts: list[str] = []
    total = 0
    for filename, data in documents or []:
        text = extract_text(filename, data)
        if not text or not text.strip():
            continue
        block = f"===== FILE: {filename} =====\n{text.strip()}" if per_file_header else text.strip()
        if max_total_chars > 0 and total + len(block) > max_total_chars:
            remaining = max_total_chars - total
            if remaining > 0:
                parts.append(block[:remaining] + "\n[...truncated...]")
            parts.append("[...remaining documents truncated to fit context...]")
            break
        parts.append(block)
        total += len(block)
    return "\n\n".join(parts)
